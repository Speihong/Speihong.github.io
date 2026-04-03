from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

doc = Document()

# 标题
title = doc.add_heading('用Trae制作汇报文档指南', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 副标题
subtitle = doc.add_paragraph('（以转正报告为例）')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 一、整体架构
doc.add_heading('一、整体架构', level=1)
doc.add_paragraph('''本项目采用HTML+Tailwind CSS实现类似PPT的演示效果，核心特点：
• 单页切换：通过JavaScript控制页面切换动画
• 响应式布局：基于Tailwind CSS的响应式设计
• 交互效果：鼠标悬停弹窗、礼花特效等''')

# 二、技术实现
doc.add_heading('二、技术实现', level=1)

doc.add_heading('1. 页面结构', level=2)
doc.add_paragraph('每个幻灯片使用<div class="slide">标签包裹，通过data-index属性标识页码。')

doc.add_heading('2. 切换动画（CSS）', level=2)
css_code = '''.slide {
    position: absolute;
    top: 0;
    left: 0;
    width: 100%;
    height: 100%;
    opacity: 0;
    transform: scale(0.98) translateX(50px);
    transition: all 0.5s cubic-bezier(0.25, 0.46, 0.45, 0.94);
    pointer-events: none;
    z-index: 1;
}
.slide.active {
    opacity: 1;
    transform: scale(1) translateX(0);
    pointer-events: auto;
    z-index: 10;
}'''
p = doc.add_paragraph(css_code)
p.runs[0].font.name = 'Consolas'

doc.add_heading('3. 翻页功能（JavaScript）', level=2)
js_code = '''// 翻页逻辑
const slides = document.querySelectorAll('.slide');
let currentSlide = 1;

function updateSlide() {
    slides.forEach((slide, index) => {
        const slideNum = index + 1;
        slide.classList.remove('active', 'prev', 'next');
        if (slideNum === currentSlide) {
            slide.classList.add('active');
        }
    });
}

// 键盘导航
document.addEventListener('keydown', (e) => {
    if (e.key === 'ArrowRight' || e.key === ' ') {
        currentSlide++;
        updateSlide();
    } else if (e.key === 'ArrowLeft') {
        currentSlide--;
        updateSlide();
    }
});'''
p = doc.add_paragraph(js_code)
p.runs[0].font.name = 'Consolas'

# 三、本地图片关联
doc.add_heading('三、本地图片关联', level=1)

doc.add_heading('1. 图片文件夹', level=2)
doc.add_paragraph('''在项目根目录创建images文件夹，将本地图片放入：

project/
├── images/
│   ├── image1.png      # 背景图
│   ├── image2.png      # Logo
│   ├── 1.png          # 展示图片1
│   └── logo.png       # 公司Logo
├── index.html
└── server.js''')

doc.add_heading('2. HTML中引用本地图片', level=2)
doc.add_paragraph('''<!-- 背景图片 -->
<div style="background-image: url('images/image1.png');">

<!-- Logo图片 -->
<img src="images/image2.png" alt="logo">

<!-- 本地图片 -->
<img src="images/1.png" alt="图片1">
''')

doc.add_heading('3. 注意事项', level=2)
doc.add_paragraph('''• 图片路径使用相对路径（相对于HTML文件）
• 中文文件名可能存在编码问题，建议使用英文命名
• 示例：工单截图.png → work_order.png''')

# 四、交互效果详解
doc.add_heading('四、交互效果详解', level=1)

doc.add_heading('1. 鼠标悬停弹窗', level=2)
doc.add_paragraph('''<!-- HTML结构 -->
<div onmouseenter="showTooltip(this)" onmouseleave="hideTooltip(this)">
    <span>悬停触发元素</span>
    <div class="tooltip-content">
        <img src="images/5.png">
    </div>
</div>

<!-- JavaScript -->
function showTooltip(element) {
    const tooltip = element.querySelector('.tooltip-content');
    tooltip.style.opacity = '1';
}

function hideTooltip(element) {
    const tooltip = element.querySelector('.tooltip-content');
    tooltip.style.opacity = '0';
}''')

doc.add_heading('2. 弹窗定位技巧', level=2)
doc.add_paragraph('''• 固定定位居中：fixed left-1/2 top-1/2 -translate-x-1/2 -translate-y-1/2
• 左侧对齐：fixed left-0 top-1/2 -translate-y-1/2
• 宽度设置：style="width: 90vw; max-width: 1200px;"''')

doc.add_heading('3. 礼花特效', level=2)
doc.add_paragraph('''// 检测当前页并触发
function checkSlideForConfetti() {
    if (currentSlide === 9) {
        setTimeout(() => {
            createConfetti();
        }, 3000);
    }
}''')

# 五、与Trae对话的技巧
doc.add_heading('五、与Trae对话的技巧', level=1)

doc.add_heading('1. 明确指定位置', level=2)
doc.add_paragraph('''• 使用"第X页"定位具体页面
• 引用HTML元素类型：div、span、p、h2等
• 描述具体内容帮助定位''')

doc.add_heading('2. 精确描述需求', level=2)
doc.add_paragraph('''✓ "将第5页的h2标题文字改为：价值反哺"
✓ "在第7页的优势卡片添加鼠标悬停弹窗"
✗ "帮我改一下那个标题"''')

doc.add_heading('3. 分步骤操作', level=2)
doc.add_paragraph('''• 大改动先备份：cp index.html index_backup.html
• 复杂效果先确认实现方式
• 完成后及时预览验证''')

doc.add_heading('4. 图片处理', level=2)
doc.add_paragraph('''• 先确认图片已放入images文件夹
• 中文文件名可能需要改为英文
• 指定图片路径时保持一致性''')

# 六、项目文件说明
doc.add_heading('六、项目文件说明', level=1)
doc.add_paragraph('''index_review_v3.html    # 主文件（当前版本）
index_review_v2.html    # 备份版本2
slides_v3/             # 截图图片文件夹
转正汇报_v3.pptx       # 导出的PPT
server.js              # 本地服务器
capture_v3.js          # 截图脚本
create_ppt.py          # PPT生成脚本''')

# 七、快速开始
doc.add_heading('七、快速开始', level=1)

doc.add_heading('1. 启动本地服务器', level=2)
doc.add_paragraph('node server.js\n# 访问 http://127.0.0.1:8082')

doc.add_heading('2. 修改内容', level=2)
doc.add_paragraph('用VS Code打开index_review_v3.html，找到对应元素进行修改')

doc.add_heading('3. 导出PPT', level=2)
doc.add_paragraph('''# 1. 先启动服务器
node server.js

# 2. 运行截图
node capture_v3.js

# 3. 生成PPT
py create_ppt.py''')

# 八、常见问题
doc.add_heading('八、常见问题', level=1)

doc.add_paragraph('''Q: 弹窗显示不出来？
A: 检查：1) 图片路径是否正确 2) JavaScript是否正确引入 3) 元素是否有relative/fixed定位

Q: 图片加载失败？
A: 1) 确认images文件夹存在 2) 检查文件名大小写 3) 尝试改为英文名

Q: 动画效果不流畅？
A: 检查CSS中transition属性设置，确保hardware acceleration

Q: 如何添加新页面？
A: 在最后一个slide后添加新的div，data-index递增''')

# 九、进阶技巧
doc.add_heading('九、进阶技巧', level=1)

doc.add_heading('1. 多种交互效果', level=2)
doc.add_paragraph('''• 淡入淡出：opacity + transition
• 缩放效果：transform: scale()
• 位移效果：transform: translateX/Y
• 旋转效果：transform: rotate()''')

doc.add_heading('2. 响应式适配', level=2)
doc.add_paragraph('''• 移动端：添加额外的媒体查询
• 平板：根据屏幕宽度调整布局
• 保持比例：使用vw/vh单位''')

doc.add_heading('3. 性能优化', level=2)
doc.add_paragraph('''• 图片压缩后再使用
• 避免过多的DOM操作
• 使用CSS transform而非left/top动画''')

# 总结
doc.add_heading('总结', level=1)
doc.add_paragraph('''通过本项目实践，我们展示了：
1. 用HTML+CSS+JS实现PPT演示效果
2. 本地资源的正确引用方式
3. 与AI协作完成复杂项目的流程

核心要点：清晰的需求描述 + 及时的效果验证 + 适当的备份策略''')

# 保存
doc.save('Trae制作汇报文档指南.docx')
print('Word文档已生成: Trae制作汇报文档指南.docx')
